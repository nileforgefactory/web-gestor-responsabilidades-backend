"""Exporta una Ficha MGA a un documento Word (.docx) descargable."""

from __future__ import annotations

import io

from docx import Document

from app.slices.sgr.models import FichaMGA

_DISCLAIMER = (
    "Documento generado automáticamente como insumo preliminar. Debe ser "
    "revisado y validado por el equipo formulador antes de su radicación "
    "oficial ante el DNP."
)


def generar_docx_ficha(*, ficha: FichaMGA, proyecto_nombre: str) -> bytes:
    doc = Document()
    doc.add_heading("Propuesta de Ficha MGA (borrador preliminar)", level=1)
    doc.add_paragraph(proyecto_nombre).italic = True
    disclaimer = doc.add_paragraph(_DISCLAIMER)
    disclaimer.italic = True

    secciones = [
        ("1. Identificación", ficha.identificacion),
        ("2. Preparación", ficha.preparacion),
        ("3. Evaluación", ficha.evaluacion),
        ("4. Programación", ficha.programacion),
    ]
    for titulo, contenido in secciones:
        doc.add_heading(titulo, level=2)
        doc.add_paragraph(contenido or "(Sección no generada aún)")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
